import time

import pytest
from docx import Document

from jobfinder import config
from jobfinder.db.models import ResumeVariant
from jobfinder.resume import cache
from jobfinder.resume.parser import parse_resume


def _build_synthetic_resume(path, skills_cell="Python, SQL, Docker"):
    """Build a small synthetic .docx exercising each heading-detection path.

    This is fixture data only - not a real resume - used solely to validate
    the parser's heuristics.
    """
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("jane@example.com")

    # All-caps heading (no Word heading style).
    document.add_paragraph("EXPERIENCE")
    document.add_paragraph("- Built a widget pipeline at Acme")
    document.add_paragraph("\u2022 Led a team of 3 engineers")

    # Word "Heading" style.
    document.add_heading("Education", level=2)
    document.add_paragraph("BS Computer Science, State University")

    # Bold short line without a heading style or all-caps text.
    bold_para = document.add_paragraph()
    bold_run = bold_para.add_run("Certifications")
    bold_run.bold = True
    document.add_paragraph("AWS Certified Developer")

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Skills"
    table.rows[0].cells[1].text = skills_cell

    document.save(str(path))


def test_parse_resume_detects_sections_and_bullets(tmp_path):
    resume_file = tmp_path / "sample.docx"
    _build_synthetic_resume(resume_file)

    data = parse_resume(resume_file)

    heading_names = [s.heading for s in data.sections]
    assert "EXPERIENCE" in heading_names
    assert "Education" in heading_names
    assert "Certifications" in heading_names

    experience = next(s for s in data.sections if s.heading == "EXPERIENCE")
    assert "Built a widget pipeline at Acme" in experience.lines
    assert "Led a team of 3 engineers" in experience.lines
    # Bullet prefixes should be stripped, not left in the cleaned lines.
    assert not any(line.startswith(("-", "\u2022")) for line in experience.lines)

    assert "Python, SQL, Docker" in data.raw_text


def test_parse_resume_missing_file_raises(tmp_path):
    with pytest.raises(FileNotFoundError):
        parse_resume(tmp_path / "does_not_exist.docx")


def test_parse_resume_empty_document_raises(tmp_path):
    empty = tmp_path / "empty.docx"
    Document().save(str(empty))
    with pytest.raises(ValueError):
        parse_resume(empty)


def test_get_or_parse_resume_caches_and_invalidates_on_change(tmp_path, monkeypatch, mocker):
    resumes_dir = tmp_path / "resumes"
    cache_dir = tmp_path / "cache"
    resumes_dir.mkdir()
    monkeypatch.setattr(config, "RESUME_DIR", resumes_dir)
    monkeypatch.setattr(config, "RESUME_CACHE_DIR", cache_dir)

    resume_file = cache.resume_path(ResumeVariant.FULLSTACK)
    _build_synthetic_resume(resume_file)

    spy = mocker.spy(cache, "parse_resume")

    data1 = cache.get_or_parse_resume(ResumeVariant.FULLSTACK)
    assert spy.call_count == 1
    assert cache._cache_path(ResumeVariant.FULLSTACK).exists()

    data2 = cache.get_or_parse_resume(ResumeVariant.FULLSTACK)
    assert spy.call_count == 1  # served from cache, no reparse
    assert data1 == data2

    # Touch the source with new content/mtime -> cache must invalidate.
    time.sleep(0.01)
    _build_synthetic_resume(resume_file, skills_cell="Go, Kubernetes")
    data3 = cache.get_or_parse_resume(ResumeVariant.FULLSTACK)
    assert spy.call_count == 2
    assert "Go, Kubernetes" in data3["raw_text"]


def test_get_or_parse_resume_missing_file_message(tmp_path, monkeypatch):
    monkeypatch.setattr(config, "RESUME_DIR", tmp_path / "resumes")
    monkeypatch.setattr(config, "RESUME_CACHE_DIR", tmp_path / "cache")
    with pytest.raises(FileNotFoundError, match="Expected resume at"):
        cache.get_or_parse_resume(ResumeVariant.PROJECT_MANAGER)
