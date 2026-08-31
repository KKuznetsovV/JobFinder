"""Parse a .docx resume into a structured, section-based representation.

The exact resume format isn't known ahead of time, so this uses conservative
heuristics (Word heading styles, bold/all-caps lines, bullet paragraphs)
rather than assuming a fixed schema. Downstream consumers (cover-letter
generator, resume selector) should treat the parsed sections as the grounded
source of truth and never invent resume content beyond what's here.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

_HEADING_STYLE_RE = re.compile(r"^heading", re.IGNORECASE)
_ALLCAPS_HEADING_RE = re.compile(r"^[A-Z0-9][A-Z0-9 &/\-]{2,40}$")
_BULLET_PREFIX_RE = re.compile(r"^[\u2022\u25CF\u25AA\-\*]\s+")


@dataclass
class ResumeSection:
    heading: str
    lines: list[str] = field(default_factory=list)


@dataclass
class ResumeData:
    source_path: str
    raw_text: str
    sections: list[ResumeSection]

    def as_dict(self) -> dict:
        return {
            "source_path": self.source_path,
            "raw_text": self.raw_text,
            "sections": [{"heading": s.heading, "lines": s.lines} for s in self.sections],
        }


def _is_heading(paragraph) -> bool:
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    if _HEADING_STYLE_RE.match(style_name):
        return True
    text = paragraph.text.strip()
    if not text or len(text) > 40:
        return False
    if _ALLCAPS_HEADING_RE.match(text) and text.upper() == text:
        return True
    # A short line where every run is bold reads as a heading even without a
    # Word heading style or all-caps formatting.
    runs = [r for r in paragraph.runs if r.text.strip()]
    if runs and all(r.bold for r in runs) and len(text.split()) <= 8:
        return True
    return False


def _clean_line(text: str) -> str:
    return _BULLET_PREFIX_RE.sub("", text.strip()).strip()


def parse_resume(path: Path | str) -> ResumeData:
    """Parse a .docx resume into raw text plus best-effort sections.

    Raises FileNotFoundError if the file doesn't exist, and ValueError if it
    has no extractable text (likely a corrupt or non-resume .docx file).
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {path}")

    document = Document(str(path))
    preamble = ResumeSection(heading="_preamble")
    sections: list[ResumeSection] = [preamble]
    current = preamble
    raw_lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        raw_lines.append(text)
        if _is_heading(paragraph):
            current = ResumeSection(heading=text)
            sections.append(current)
        else:
            current.lines.append(_clean_line(text))

    # Many resumes put the skills/contact block in a table rather than
    # paragraphs; fold that text into whichever section came last.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    raw_lines.append(cell_text)
                    current.lines.append(_clean_line(cell_text))

    raw_text = "\n".join(raw_lines)
    if not raw_text:
        raise ValueError(f"No extractable text found in resume: {path}")

    sections = [s for s in sections if s.lines or s.heading != "_preamble"]
    return ResumeData(source_path=str(path), raw_text=raw_text, sections=sections)
